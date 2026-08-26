import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title_slide(doc, title, subtitle):
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(subtitle)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

def add_slide(doc, title, bullet_points):
    doc.add_heading(title, level=1)
    for point in bullet_points:
        doc.add_paragraph(point, style='List Bullet')
    doc.add_page_break()

def add_weak_topic_slide(doc):
    doc.add_heading("5. Advanced Logic: Weak Topic & Subject Classification", level=1)
    
    p1 = doc.add_paragraph("The platform employs a robust algorithmic model to identify a student's academic weaknesses at the microscopic (topic) level. It evaluates the ratio of 'non-positive' attempts (incorrect answers + unattempted questions) against the total questions for a specific topic.")
    
    doc.add_heading("Classification Thresholds:", level=2)
    doc.add_paragraph("Strong Weakness ('Weakest Topic'):", style='List Bullet')
    doc.add_paragraph("Triggered when (Incorrect + Unattempted) / Total Questions ≥ 66.6% (2/3).", style='List Continue')
    doc.add_paragraph("Indicates severe conceptual gaps requiring immediate intervention.", style='List Continue')
    
    doc.add_paragraph("Medium Weakness ('Weak Topic'):", style='List Bullet')
    doc.add_paragraph("Triggered when (Incorrect + Unattempted) / Total Questions is between 50% (1/2) and 66.6%.", style='List Continue')
    doc.add_paragraph("Indicates partial understanding or frequent silly mistakes.", style='List Continue')
    
    doc.add_heading("Database Architecture:", level=2)
    doc.add_paragraph("Stored via the StudentWeakTopics MongoDB Schema.", style='List Bullet')
    doc.add_paragraph("Data is aggregated hierarchically: Questions → Topics → Subjects (Physics, Chemistry, Math).", style='List Bullet')
    doc.add_paragraph("Enables Center Managers to see exactly which chapters (e.g., Thermodynamics) are causing the most score bleed.", style='List Bullet')
    
    doc.add_page_break()

doc = Document()

# Title Slide
add_title_slide(
    doc,
    "Pragati\nPerformance & Analytics Platform",
    "A Comprehensive Technical & Analytical Overview"
)

# Slide 1: Introduction
add_slide(
    doc,
    "1. Introduction to Pragati",
    [
        "Pragati is a dedicated data analytics and management platform tailored for CSRL.",
        "Designed to track student performance, analyze center-level progress, and monitor subjective trends across various FMT tests.",
        "Provides actionable insights through dynamically rendered Top/Bottom rankings, performance graphs, and weak topic analysis.",
        "Equips administrators and center managers with a unified, real-time reporting interface."
    ]
)

# Slide 2: Technology Stack
add_slide(
    doc,
    "2. Core Technology Stack",
    [
        "Frontend (UI/UX): React.js utilizing Vite/Create-React-App for rapid compilation and component reusability.",
        "Styling: Pure custom CSS with modern design paradigms (Glassmorphism, CSS variables) for a polished, seamless experience.",
        "Data Visualization: Recharts library used for dynamic plotting of line charts and bar charts across multiple filters.",
        "Backend Architecture: Node.js with Express framework acting as a robust RESTful API layer.",
        "Database: MongoDB (Mongoose) storing complex, hierarchical test records, profiles, and analytical datasets.",
        "Deployment & CI/CD: Vercel for instant frontend hosting and edge caching."
    ]
)

# Slide 3: Qualification Logic
add_slide(
    doc,
    "3. Core Logic: Qualification Metrics",
    [
        "Subject-wise Qualification Threshold: A student is considered 'Qualified' ONLY if they score ≥ 30 in Physics, ≥ 30 in Chemistry, AND ≥ 30 in Mathematics/Biology.",
        "Single Test Appearance: For an individual test, the 'Appeared' count is the exact number of students who submitted scores.",
        "Multiple Test Selection (e.g., ALL FMT): To prevent inflated baseline figures, the 'Appeared' count is calculated as the Average number of appearances across all selected tests.",
        "Aggregate Qualification: The 'Qualified' count for multiple tests is the RAW SUM of all qualifying instances across those tests.",
        "Percentage Formula: (Raw Qualified Sum / Average Appeared) * 100. This provides a realistic success density without skewing the denominator."
    ]
)

# Slide 4: Ranking & Leaderboard Logic
add_slide(
    doc,
    "4. Core Logic: Rankings & Leaderboards",
    [
        "Dynamic Aggregation: Students are sorted in descending order by their Total Marks for the active test filter.",
        "Tie-breaker: If marks are identical, ties are resolved alphabetically by the student's Roll Number.",
        "Historical Traceability: The leaderboard explicitly injects historical ranking columns (e.g., FMT07 Rank, FMT06 Rank) directly alongside the current score.",
        "Absentee Handling: Students who missed a test or lack a score in the DB are systematically flagged as 'Absent' rather than a null dash.",
        "Dual Views: Dedicated panels restrict top-performers (Top 15) and under-performers (Bottom 15) to maintain concise, readable interfaces.",
        "Column Order Priority: Subject marks (C, M, P) are displayed sequentially, followed immediately by the Total, keeping scoring contexts tight."
    ]
)

# Slide 5: Weak Topic Logic
add_weak_topic_slide(doc)

# Slide 6: Future Scalability
add_slide(
    doc,
    "6. Future Roadmap & Scalability",
    [
        "Modular Codebase: Fully componentized React architecture allows drop-in additions of new features (e.g., red flag visualizers for subject averages ≤ 20).",
        "PDF/Excel Exporting: Capable of generating printable PDF reports and data-sheets directly from the DOM.",
        "Predictive Modeling: Data structures are primed to incorporate ML-based predictions on student trajectories based on historical weak topics."
    ]
)

output_path = '/Users/surya/Desktop/Pragati_Presentation.docx'
doc.save(output_path)
print(f"Presentation saved to {output_path}")

