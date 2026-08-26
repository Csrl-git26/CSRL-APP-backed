import sys
from docx import Document
from docx.shared import Pt, Inches

doc = Document()

# Add Subject
subject_p = doc.add_paragraph()
subject_p.add_run("Subject: ").bold = True
subject_p.add_run("Introduction of प्रगति Dashboard – Web-Based Student Performance Analytics Platform")

doc.add_paragraph("Dear Centre Managers & Faculty,")
doc.add_paragraph("We are pleased to introduce प्रगति Dashboard — CSRL's newly developed web-based student performance analytics and data management platform, designed to empower our faculties with data-driven insights and bring greater transparency to student tracking.")

# Section A
p_A = doc.add_paragraph()
p_A.add_run("A) What is प्रगति Dashboard?").bold = True
doc.add_paragraph("प्रगति is a centralized web-based platform through which student academic performance can be tracked and analysed across your centre. It provides a single interface for monitoring test performance, subject-wise rankings, student progress, and centre-level academic performance.")

# Section B
p_B = doc.add_paragraph()
p_B.add_run("B) Phase-01 — Current Status").bold = True
doc.add_paragraph("As part of Phase-01, we have currently uploaded previous session data points on the platform. This has been done to test the system, understand the complete process, validate the analysis and familiarize the faculty and centre teams with the dashboard.")
doc.add_paragraph("In the coming week, we will upload the CMT-01, MT-01 and MT-02 data for the 2026-27 session. The platform will then progressively move towards regular use with current-session data.")

# Section C
p_C = doc.add_paragraph()
p_C.add_run("C) Key Features for Centre Management & Faculty").bold = True
doc.add_paragraph("Access your centre-level performance metrics instantly.", style='List Bullet')
doc.add_paragraph("Identify weak subjects and subject-wise weak topics across your student group to target interventions.", style='List Bullet')
doc.add_paragraph("Generate student performance reports that can be exported and shared with parents for regular academic tracking and discussion.", style='List Bullet')
doc.add_paragraph("Track individual student performance, progress trends, and areas requiring immediate academic support.", style='List Bullet')
doc.add_paragraph("Review complete test records, attendance, and performance trends for timely intervention.", style='List Bullet')

# Section D
p_D = doc.add_paragraph()
p_D.add_run("D) Login Details & Access").bold = True
doc.add_paragraph("Platform Link: https://csrl-pragati.vercel.app", style='List Bullet')
doc.add_paragraph("Role Selection: On the login page, please select the \"Centre Faculty\" tab.", style='List Bullet')
doc.add_paragraph("Username: [Your Assigned Centre Username] (e.g. KNP, DDN, JDH, etc.)", style='List Bullet')
doc.add_paragraph("Password: [Your Assigned Centre Password]", style='List Bullet')

# Section E
p_E = doc.add_paragraph()
p_E.add_run("E) Tangible Outputs & Impact on the Centre").bold = True

p_E_outputs = doc.add_paragraph()
p_E_outputs.add_run("Outputs You Will Receive:").bold = True
doc.add_paragraph("Automated Student Report Cards: Instantly generate downloadable PDF performance reports for individual students.", style='List Bullet')
doc.add_paragraph("Weak Topic Analytics: Get exact lists of chapters and subjects where your specific batch is losing marks.", style='List Bullet')
doc.add_paragraph("Progress Trend Graphs: Visual charts tracking your students' attendance, accuracy, and marks over the entire academic year.", style='List Bullet')

p_E_impact = doc.add_paragraph()
p_E_impact.add_run("Effect & Impact After Using the Platform:").bold = True
doc.add_paragraph("Targeted Teaching: Faculty will no longer need to guess where students are struggling; they can dedicate revision classes specifically to the 'Weak Topics' flagged by the system.", style='List Bullet')
doc.add_paragraph("Time Savings: Eliminates the manual administrative effort of calculating averages, subject cutoffs, and ranks for hundreds of students.", style='List Bullet')
doc.add_paragraph("Stronger Parent Engagement: Instantly sharing structured, data-backed reports builds trust and keeps parents deeply involved in their child's academic journey.", style='List Bullet')
doc.add_paragraph("Improved Centre Results: By identifying struggling students early and intervening before final exams, the overall qualification rate and rank density of the centre will significantly improve.", style='List Bullet')

# Section F
p_F = doc.add_paragraph()
p_F.add_run("F) Why This Matters").bold = True
doc.add_paragraph("Identify struggling students early and enable proactive academic intervention.", style='List Bullet')
doc.add_paragraph("Support data-backed decisions on curriculum focus.", style='List Bullet')
doc.add_paragraph("Build accountability at the student and faculty levels.", style='List Bullet')

# Outro
doc.add_paragraph()
doc.add_paragraph("A special acknowledgement to Mr. Ajay, who has made a significant contribution towards the development and implementation of the प्रगति Dashboard. We also appreciate the efforts of the team members who have been involved in testing, validation and providing feedback at different stages of the platform.")
doc.add_paragraph("The development of this platform has been a collaborative effort, and the contribution of everyone involved is appreciated.")
doc.add_paragraph("We believe प्रगति Dashboard will become an important tool for strengthening academic monitoring, improving intervention and supporting better student outcomes at your centre.")

p_req = doc.add_paragraph()
p_req_run = p_req.add_run("We request all centre faculty and managers to explore the platform and share their feedback and suggestions for further improvement by 30th August 2026.")
p_req_run.bold = True
p_req_run.italic = True

doc.add_paragraph()
doc.add_paragraph("Best Regards,")
p_sig = doc.add_paragraph()
p_sig.add_run("ACADEMICS DEPARTMENT-CSRL").bold = True

output_path = '/Users/surya/Desktop/Centre_Faculty_Email_Draft.docx'
doc.save(output_path)
print(f"Saved to {output_path}")
